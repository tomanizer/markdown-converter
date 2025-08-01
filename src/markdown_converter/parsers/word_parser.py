"""
Word Document Parser

This module provides specialized parsing for Word documents (.docx, .doc)
using mammoth as the primary processor and python-docx as a fallback.
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional, List, Union
import tempfile
import os

from .base import BaseParser, ParserResult
from ..core.exceptions import ParserError, UnsupportedFormatError


class WordParser(BaseParser):
    """
    Specialized parser for Word documents (.docx, .doc).
    
    Uses mammoth as the primary processor for clean HTML conversion,
    with python-docx as a fallback for complex documents.
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None) -> None:
        """
        Initialize the Word parser.
        
        :param config: Configuration dictionary
        """
        super().__init__(config)
        self.logger = logging.getLogger("WordParser")
        self._setup_default_config()
        self._validate_dependencies()
    
    def _setup_default_config(self) -> None:
        """Setup default configuration for Word parsing."""
        self.default_config = {
            # Primary processor (mammoth)
            "use_mammoth": True,
            "mammoth_options": {
                "convert_image": True,
                "ignore_empty_paragraphs": True,
                "id_prefix": "docx-",
                "transform_document": None,
            },
            # Fallback processor (python-docx)
            "use_python_docx": True,
            "python_docx_options": {
                "extract_images": True,
                "preserve_formatting": True,
                "include_metadata": True,
            },
            # Output settings
            "extract_images": True,
            "image_dir": "images",
            "preserve_tables": True,
            "preserve_headers": True,
            "preserve_footers": True,
        }
        
        # Merge with user config
        if self.config:
            self.default_config.update(self.config)
    
    def _validate_dependencies(self) -> None:
        """Validate that required dependencies are available."""
        try:
            import mammoth
            self.mammoth_available = True
        except ImportError:
            self.mammoth_available = False
            self.logger.warning("mammoth not available, will use python-docx fallback")
        
        try:
            import docx
            self.python_docx_available = True
        except ImportError:
            self.python_docx_available = False
            self.logger.warning("python-docx not available")
        
        if not self.mammoth_available and not self.python_docx_available:
            raise ParserError("Neither mammoth nor python-docx are available")
    
    def can_parse(self, file_path: Union[str, Path]) -> bool:
        """
        Check if this parser can handle the given file.
        
        :param file_path: Path to the file
        :return: True if the file can be parsed
        """
        file_path = Path(file_path)
        return file_path.suffix.lower() in ['.docx', '.doc']
    
    def parse(self, file_path: Union[str, Path]) -> ParserResult:
        """
        Parse a Word document and extract its content.
        
        :param file_path: Path to the Word document
        :return: ParserResult with extracted content and metadata
        :raises: ParserError if parsing fails
        """
        file_path = Path(file_path)
        
        if not self.can_parse(file_path):
            raise UnsupportedFormatError(f"Cannot parse {file_path.suffix} files")
        
        self.logger.info(f"Parsing Word document: {file_path}")
        
        try:
            # Try mammoth first (preferred)
            if self.default_config["use_mammoth"] and self.mammoth_available:
                try:
                    return self._parse_with_mammoth(file_path)
                except Exception as e:
                    self.logger.warning(f"Mammoth parsing failed: {e}, trying python-docx")
            
            # Fallback to python-docx
            if self.default_config["use_python_docx"] and self.python_docx_available:
                return self._parse_with_python_docx(file_path)
            
            raise ParserError("No available Word parser found")
            
        except Exception as e:
            self.logger.error(f"Word parsing failed for {file_path}: {e}")
            raise ParserError(f"Failed to parse Word document {file_path}: {e}")
    
    def _parse_with_mammoth(self, file_path: Path) -> ParserResult:
        """
        Parse Word document using mammoth.
        
        :param file_path: Path to the Word document
        :return: ParserResult with extracted content
        """
        import mammoth
        
        self.logger.debug(f"Using mammoth to parse {file_path}")
        
        # Setup mammoth options
        options = self.default_config["mammoth_options"].copy()
        
        # Convert to HTML first
        with open(file_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file, options)
        
        # Extract content and metadata
        content = result.value
        messages = result.messages
        
        # Process any warnings
        if messages:
            self.logger.warning(f"Mammoth messages for {file_path}: {messages}")
        
        # Extract metadata
        metadata = self._extract_metadata_mammoth(file_path)
        
        return ParserResult(
            content=content,
            metadata=metadata,
            format="html",  # Mammoth outputs HTML
            messages=messages
        )
    
    def _parse_with_python_docx(self, file_path: Path) -> ParserResult:
        """
        Parse Word document using python-docx.
        
        :param file_path: Path to the Word document
        :return: ParserResult with extracted content
        """
        from docx import Document
        
        self.logger.debug(f"Using python-docx to parse {file_path}")
        
        # Load the document
        doc = Document(file_path)
        
        # Extract content
        content = self._extract_content_python_docx(doc)
        
        # Extract metadata
        metadata = self._extract_metadata_python_docx(doc)
        
        return ParserResult(
            content=content,
            metadata=metadata,
            format="markdown",
            messages=[]
        )
    
    def _extract_content_python_docx(self, doc) -> str:
        """
        Extract content from python-docx document.
        
        :param doc: python-docx Document object
        :return: Extracted content as markdown
        """
        content_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Handle different paragraph styles
                style = paragraph.style.name.lower()
                text = paragraph.text.strip()
                
                if 'heading' in style:
                    level = style.replace('heading', '').strip()
                    if level.isdigit():
                        heading_level = int(level)
                    else:
                        heading_level = 1
                    content_parts.append(f"{'#' * heading_level} {text}")
                else:
                    content_parts.append(text)
                
                content_parts.append("")  # Add blank line
        
        # Extract tables
        for table in doc.tables:
            content_parts.append(self._convert_table_to_markdown(table))
            content_parts.append("")  # Add blank line
        
        return "\n".join(content_parts)
    
    def _convert_table_to_markdown(self, table) -> str:
        """
        Convert a python-docx table to markdown.
        
        :param table: python-docx table object
        :return: Markdown table string
        """
        if not table.rows:
            return ""
        
        markdown_lines = []
        
        # Process each row
        for i, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                # Get cell text and clean it
                cell_text = cell.text.strip().replace('\n', ' ')
                cells.append(cell_text)
            
            # Create markdown row
            markdown_row = "| " + " | ".join(cells) + " |"
            markdown_lines.append(markdown_row)
            
            # Add separator row after header
            if i == 0:
                separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                markdown_lines.append(separator)
        
        return "\n".join(markdown_lines)
    
    def _extract_metadata_mammoth(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata using mammoth.
        
        :param file_path: Path to the Word document
        :return: Dictionary of metadata
        """
        # Mammoth doesn't provide much metadata, so we'll use file info
        return {
            "parser": "mammoth",
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size,
            "format": "docx"
        }
    
    def _extract_metadata_python_docx(self, doc) -> Dict[str, Any]:
        """
        Extract metadata from python-docx document.
        
        :param doc: python-docx Document object
        :return: Dictionary of metadata
        """
        metadata = {
            "parser": "python-docx",
            "format": "docx"
        }
        
        # Extract core properties if available
        try:
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
        except Exception as e:
            self.logger.warning(f"Could not extract core properties: {e}")
        
        # Extract document statistics
        try:
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            metadata["section_count"] = len(doc.sections)
        except Exception as e:
            self.logger.warning(f"Could not extract document statistics: {e}")
        
        return metadata
    
    def get_supported_formats(self) -> List[str]:
        """
        Get list of supported file formats.
        
        :return: List of supported format extensions
        """
        return ['.docx', '.doc']
    
    def get_parser_info(self) -> Dict[str, Any]:
        """
        Get information about this parser.
        
        :return: Dictionary with parser information
        """
        return {
            "name": "WordParser",
            "description": "Specialized parser for Word documents",
            "supported_formats": self.get_supported_formats(),
            "dependencies": {
                "mammoth": self.mammoth_available,
                "python-docx": self.python_docx_available
            },
            "config": self.default_config
        } 