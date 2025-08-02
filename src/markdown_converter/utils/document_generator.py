"""
Document Generator for Test Files

This module provides utilities to create test documents in various formats
for testing the markdown converter functionality.
"""

import json
import os
from pathlib import Path
from typing import Any, Dict, List, Union

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl import Workbook

    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


class DocumentGenerator:
    """Generate test documents in various formats."""

    def __init__(self) -> None:
        """Initialize the document generator."""
        self.supported_formats = {
            "word": self._create_word_document,
            "excel": self._create_excel_document,
            "html": self._create_html_document,
            "outlook": self._create_outlook_document,
            "text": self._create_text_document,
        }

    def create_document(
        self, doc_type: str, content: Dict[str, Any], output_path: Path
    ) -> None:
        """
        Create a document of the specified type.

        :param doc_type: Type of document to create ('word', 'excel', 'html', etc.)
        :param content: Content dictionary for the document
        :param output_path: Path where to save the document
        """
        if doc_type not in self.supported_formats:
            raise ValueError(f"Unsupported document type: {doc_type}")

        creator = self.supported_formats[doc_type]
        creator(content, output_path)

    def _create_word_document(self, content: Dict[str, Any], output_path: Path) -> None:
        """Create a Word document."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not available")

        doc = Document()

        # Add title
        if "title" in content:
            title = doc.add_heading(content["title"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add sections
        if "sections" in content:
            for section in content["sections"]:
                if "heading" in section:
                    doc.add_heading(section["heading"], level=1)

                if "content" in section:
                    doc.add_paragraph(section["content"])

                if "table" in section:
                    table_data = section["table"]
                    if table_data:
                        table = doc.add_table(
                            rows=len(table_data), cols=len(table_data[0])
                        )
                        table.style = "Table Grid"

                        for i, row in enumerate(table_data):
                            for j, cell_value in enumerate(row):
                                table.cell(i, j).text = str(cell_value)

        doc.save(str(output_path))

    def _create_excel_document(
        self, content: Dict[str, Any], output_path: Path
    ) -> None:
        """Create an Excel document."""
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl is not available")

        wb = Workbook()

        # Remove default sheet
        wb.remove(wb.active)

        # Add sheets
        if "sheets" in content:
            for sheet_info in content["sheets"]:
                ws = wb.create_sheet(title=sheet_info["name"])

                if "data" in sheet_info:
                    for row in sheet_info["data"]:
                        ws.append(row)

        wb.save(str(output_path))

    def _create_html_document(self, content: Dict[str, Any], output_path: Path) -> None:
        """Create an HTML document."""
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>{content.get('title', 'Test Document')}</title>
    <meta charset="utf-8">
</head>
<body>
{content.get('body', '<h1>Test Document</h1><p>This is a test document.</p>')}
</body>
</html>"""

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)

    def _create_outlook_document(
        self, content: Dict[str, Any], output_path: Path
    ) -> None:
        """Create an Outlook email document (.msg file)."""
        # For now, create a simple text representation
        # In a real implementation, you'd use extract-msg or similar
        msg_content = f"""Subject: {content.get('subject', 'Test Email')}
From: {content.get('from', 'test@example.com')}
To: {content.get('to', 'recipient@example.com')}
Date: {content.get('date', '2024-01-15')}

{content.get('body', 'This is a test email.')}"""

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(msg_content)

    def _create_text_document(self, content: Dict[str, Any], output_path: Path) -> None:
        """Create a plain text document."""
        text_content = content.get("content", "This is a test document.")

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text_content)

    def create_simple_test_files(self, output_dir: Path) -> None:
        """Create simple test files for basic testing."""

        # Simple HTML
        html_content = """<!DOCTYPE html>
<html>
<head><title>Simple Test</title></head>
<body>
<h1>Simple Test Document</h1>
<p>This is a simple HTML test document.</p>
<ul>
<li>First item</li>
<li>Second item</li>
<li>Third item</li>
</ul>
</body>
</html>"""

        with open(output_dir / "simple_test.html", "w") as f:
            f.write(html_content)

        # Simple text
        text_content = """Simple Test Document
====================

This is a simple text document for testing.

Features:
- Basic text
- Simple formatting
- Easy to parse

End of document."""

        with open(output_dir / "simple_test.txt", "w") as f:
            f.write(text_content)

        # Simple CSV (Excel-like)
        csv_content = """Name,Age,City
John Doe,30,New York
Jane Smith,25,Los Angeles
Bob Johnson,35,Chicago"""

        with open(output_dir / "simple_test.csv", "w") as f:
            f.write(csv_content)
